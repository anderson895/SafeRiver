"""
Builds documentation.docx from DOCUMENTATION.md, with screenshots embedded.

    npm run docs:docx

Generated rather than hand-assembled so the Word file cannot drift from the
markdown. Re-run it after editing DOCUMENTATION.md or regenerating screenshots.

Requires python-docx and Pillow, both already present.
"""

import io
import os
import re
import sys

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from PIL import Image

# One phone screenful: 390x844 CSS px at 2x device pixel ratio, matching the
# viewport scripts/screenshots.ts captures at.
MOBILE_VIEWPORT_PX = 1688

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
MD = os.path.join(ROOT, "DOCUMENTATION.md")
OUT = os.path.join(ROOT, "documentation.docx")
SHOTS = os.path.join(ROOT, "screenshot")

# Screens, in the order the documentation describes them. Desktop and mobile of
# the same screen sit together so the responsive claim is visible side by side.
FIGURES = [
    ("01-dashboard", "Dashboard — interactive flood hazard map and latest alerts"),
    ("02-rainfall", "Rainfall — radar map with time slider"),
    ("02b-rainfall-data", "Rainfall — 24-hour forecast chart"),
    ("03-water-level", "Water Level — reservoir level against Normal / Alert / Critical"),
    ("04-dam-advisory", "Dam Advisory — the Agno cascade, upstream to downstream"),
    ("05-flood-info", "Flood Information — locally written guidance"),
    ("05b-flood-info-during", "Flood Information — During a Flood"),
    ("06-alerts", "Alerts — active and past advisories"),
    ("07-subscribe", "Get Alerts — email subscription with double opt-in"),
    ("08-about", "About — data sources, limitations and privacy notice"),
]

BLUE = RGBColor(0x15, 0x65, 0xC0)
GREY = RGBColor(0x60, 0x60, 0x60)


def add_picture_fitted(doc, path, max_w_in, max_h_in=8.0):
    """Insert an image scaled to fit both a width and a height budget.

    Full-page captures are very tall; sizing on width alone pushes them past a
    single page and splits a screenshot across two.
    """
    with Image.open(path) as im:
        w, h = im.size
    ratio = h / w
    width = max_w_in
    if width * ratio > max_h_in:
        width = max_h_in / ratio
    doc.add_picture(path, width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


def mobile_image(path):
    """Return one phone screenful of a full-page mobile capture.

    Full-page mobile shots run to extreme aspect ratios — the About page is
    780x7150, over 9:1. Fitting that to the page height leaves it 0.7 in wide,
    far too small to read. Cropping to the viewport keeps every mobile figure at
    a consistent, legible 2.16:1 and shows what a phone actually displays.
    Returns (stream, was_cropped).
    """
    with Image.open(path) as im:
        w, h = im.size
        if h <= MOBILE_VIEWPORT_PX:
            buf = io.BytesIO()
            im.save(buf, format="PNG")
            buf.seek(0)
            return buf, False
        buf = io.BytesIO()
        im.crop((0, 0, w, MOBILE_VIEWPORT_PX)).save(buf, format="PNG")
        buf.seek(0)
        return buf, True


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = GREY


def parse_table(lines, i):
    """Consume a markdown table starting at line i. Returns (rows, next_i)."""
    rows = []
    while i < len(lines) and lines[i].strip().startswith("|"):
        cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
        # Skip the |---|---| separator row.
        if not all(set(c) <= set("-: ") for c in cells):
            rows.append(cells)
        i += 1
    return rows, i


def clean_inline(text):
    """Strip markdown emphasis and links; Word carries no markdown."""
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"(?<!\w)\*([^*]+)\*(?!\w)", r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    return text.strip()


def add_rich(paragraph, text):
    """Render **bold** spans as real bold runs."""
    for part in re.split(r"(\*\*[^*]+\*\*)", text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            paragraph.add_run(clean_inline(part[2:-2])).bold = True
        else:
            paragraph.add_run(clean_inline(part))


def build():
    if not os.path.exists(MD):
        sys.exit(f"Missing {MD}")

    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    # ---- Title page ----
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("SafeRiver")
    r.bold = True
    r.font.size = Pt(30)
    r.font.color.rgb = BLUE

    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = s.add_run(
        "Web-Based Flood Hazard Information and Dam Release Alert System\n"
        "for Agno River Communities in San Manuel, Pangasinan"
    )
    r.font.size = Pt(13)

    u = doc.add_paragraph()
    u.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = u.add_run("https://safe-river-san-manuel.vercel.app")
    r.font.size = Pt(10)
    r.font.color.rgb = GREY

    doc.add_page_break()

    # ---- Body, converted from the markdown ----
    with open(MD, encoding="utf-8") as f:
        lines = f.read().split("\n")

    i = 0
    in_code = False
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            in_code = not in_code
            i += 1
            continue

        if in_code:
            p = doc.add_paragraph()
            r = p.add_run(line)
            r.font.name = "Consolas"
            r.font.size = Pt(9)
            i += 1
            continue

        if stripped.startswith("|"):
            rows, i = parse_table(lines, i)
            if rows:
                table = doc.add_table(rows=0, cols=len(rows[0]))
                table.style = "Light Grid Accent 1"
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                for ri, row in enumerate(rows):
                    cells = table.add_row().cells
                    for ci, cell in enumerate(row[: len(cells)]):
                        para = cells[ci].paragraphs[0]
                        add_rich(para, cell)
                        for run in para.runs:
                            run.font.size = Pt(9)
                            if ri == 0:
                                run.bold = True
                doc.add_paragraph()
            continue

        if stripped.startswith("# "):
            i += 1
            continue  # the title page already carries this

        if stripped.startswith("## "):
            doc.add_heading(clean_inline(stripped[3:]), level=1)
        elif stripped.startswith("### "):
            doc.add_heading(clean_inline(stripped[4:]), level=2)
        elif stripped.startswith("---"):
            pass
        elif stripped.startswith("- ") or stripped.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            add_rich(p, stripped[2:])
        elif stripped:
            p = doc.add_paragraph()
            add_rich(p, stripped)

        i += 1

    # ---- Appendix: screenshots ----
    doc.add_page_break()
    doc.add_heading("Appendix A — Screens", level=1)
    doc.add_paragraph(
        "Captured automatically at 1440x900 (desktop) and 390x844 (mobile), the "
        "latter approximating an entry-level Android handset. Both English and "
        "Tagalog are generated; English is shown here."
    )

    missing = []
    fig = 1
    for name, caption in FIGURES:
        desktop = os.path.join(SHOTS, "desktop", "en", f"{name}.png")
        if not os.path.exists(desktop):
            missing.append(name)
            continue
        doc.add_heading(f"Figure {fig}. {caption}", level=2)
        add_picture_fitted(doc, desktop, max_w_in=6.3, max_h_in=7.5)
        add_caption(doc, f"Figure {fig}. {caption} (desktop)")
        doc.add_page_break()
        fig += 1

    # Mobile evidence for the responsive requirement, two per page.
    doc.add_heading("Appendix B — Mobile layout", level=1)
    doc.add_paragraph(
        "The same screens at 390 px width, evidence for the responsive interface "
        "requirement and for the accessibility evaluation. Each is cropped to one "
        "phone screenful — what a resident sees on opening the page, before "
        "scrolling."
    )

    mobile_shots = [
        (n, c) for n, c in FIGURES
        if os.path.exists(os.path.join(SHOTS, "mobile", "en", f"{n}.png"))
    ]
    for idx in range(0, len(mobile_shots), 2):
        pair = mobile_shots[idx : idx + 2]
        table = doc.add_table(rows=1, cols=len(pair))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for ci, (name, caption) in enumerate(pair):
            cell = table.rows[0].cells[ci]
            path = os.path.join(SHOTS, "mobile", "en", f"{name}.png")
            stream, cropped = mobile_image(path)
            run = cell.paragraphs[0].add_run()
            run.add_picture(stream, width=Inches(2.5))
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap = cell.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            label = caption + (" (scrolls further)" if cropped else "")
            r = cap.add_run(label)
            r.italic = True
            r.font.size = Pt(8)
            r.font.color.rgb = GREY
        doc.add_paragraph()

    doc.save(OUT)

    size_kb = os.path.getsize(OUT) / 1024
    print(f"Wrote {OUT}  ({size_kb:.0f} KB)")
    print(f"  {fig - 1} desktop figures, {len(mobile_shots)} mobile figures")
    if missing:
        print(f"  WARNING: no screenshot for {', '.join(missing)}")
        print("           run `npm run screenshots` first")


if __name__ == "__main__":
    build()
